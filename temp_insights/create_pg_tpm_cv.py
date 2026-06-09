from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = Path("OUTPUTS/cases/pg_senior_tpm_pampers_club_app_2026/MKorsikov_PG_Senior_TPM_Pampers_Club_CV.docx")


def set_run(run, size=10, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run(r, size=10.5, bold=True, color=(31, 78, 121))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_run(r, size=9.4)
    return p


def add_role(doc, title, org, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{title} | {org} | {location} | {dates}")
    set_run(r, size=10, bold=True)
    return p


def add_inline_section(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label)
    set_run(r, size=9.5, bold=True)
    r = p.add_run(text)
    set_run(r, size=9.5)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9.4)
styles["List Bullet"].font.name = "Arial"
styles["List Bullet"].font.size = Pt(9.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
r = p.add_run("MYKOLA KORSIKOV")
set_run(r, size=17, bold=True, color=(31, 78, 121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(7)
r = p.add_run("Poznan, Poland | +48 513 523 937 | nikolay.korsikov@gmail.com | linkedin.com/in/korsikov/")
set_run(r, size=9)

add_heading(doc, "Target Role")
add_inline_section(
    doc,
    "Senior Technical Product Manager - Pampers Club App: ",
    "technical product ownership, platform simplification, vendor challenge, delivery governance, architecture trade-offs, and scalable global product delivery.",
)

add_heading(doc, "Profile")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    "Technical product and transformation professional with 15+ years of project, portfolio, service, and team leadership experience, including 5 years in product management. Strong record translating business needs into practical technology roadmaps, governing delivery across vendors and internal teams, and challenging unnecessary complexity, cost, and over-engineering. Experienced in product roadmap ownership, backlog governance, requirements, release coordination, architecture review, integration discussions, risk/compliance documentation, and stakeholder alignment across global enterprise environments. Particularly strong in simplifying fragmented processes, balancing value/speed/cost trade-offs, and keeping solutions scalable, maintainable, and operationally sustainable."
)
set_run(r, size=9.4)

add_heading(doc, "Core Strengths for P&G")
for item in [
    "Technical product management: roadmap, backlog, requirements, release scope, governance forums, stakeholder communication, and delivery priorities.",
    "Platform simplification: translating business ambition into leaner, standardized, scalable, and maintainable solution choices.",
    "Vendor challenge: commercially aware review of estimates, scope, implementation choices, supportability, and value for money.",
    "Architecture awareness: API, integration, data-flow, cloud, supportability, security, and non-functional requirement discussions with technical teams.",
    "Delivery governance: planning, prioritization, dependency management, release cadence, change control, risk management, and production transition.",
    "Business-technology translation: explaining trade-offs between user value, speed, cost, risk, scalability, and long-term maintainability.",
]:
    add_bullet(doc, item)

add_heading(doc, "Experience")
add_role(doc, "Product Manager", "Haleon", "Poznan, Poland", "2022 - Present")
for item in [
    "Own global Planning, Budgeting and Controlling product vision, roadmap, priorities, and governance, aligning finance, technology, security, compliance, support, and leadership stakeholders.",
    "Write requirements, user stories, functional and non-functional requirements, use cases, acceptance criteria, and decision-ready documentation to guide delivery and reduce ambiguity.",
    "Review architecture proposals and assess technical feasibility, integration options, system dependencies, data flows, supportability, and fit with wider enterprise platforms.",
    "Manage and refine backlog priorities across features, defects, technical debt, stakeholder demand, operational risk, and delivery capacity.",
    "Analyze product usage, API metrics, latency, errors, KPIs, test results, and service signals to monitor performance and inform product decisions.",
    "Coordinate application integrations, interface/API discussions, ETL pipeline work with data engineers, and technical constraints with delivery teams.",
    "Create technology and compliance documentation including access plans, disaster recovery plans, BPSD, governance plans, change materials, and production readiness evidence.",
    "Plan and manage releases through ServiceNow, defining release scope, coordinating cross-team dependencies, establishing cadence, and supporting smooth production transitions.",
    "Support large-scale SaaS-enabled transformation through OnePlan solution implementation with Board SaaS at the core and total program cost above GBP 10m.",
]:
    add_bullet(doc, item)

add_role(doc, "Portfolio Manager", "GSK", "Poznan, Poland", "2020 - 2022")
for item in [
    "Kept leadership, stakeholders, vendors, and dependent teams aligned on roadmap progress, priorities, risks, dependencies, and delivery decisions.",
    "Managed vendor governance and service reviews, including contract/SOW/change request review, purchase orders, invoices, license management, and infrastructure tracking.",
    "Challenged delivery scope, cost, estimates, and vendor proposals to support value-for-money decisions and avoid unnecessary complexity.",
    "Facilitated agile ceremonies, planning, reviews, retrospectives, and delivery coordination while supporting agile ways of working across teams.",
    "Reviewed metrics, OKR alignment, and dashboards to improve KPI tracking, portfolio visibility, and leadership decision-making.",
    "Researched market trends and competitor offerings to inform roadmap direction and business prioritization.",
]:
    add_bullet(doc, item)

add_role(doc, "Team Leader", "Fujitsu", "Lodz, Poland", "2017 - 2019")
for item in [
    "Led cross-functional operational and transformation initiatives across Finance, HR, Marketing, and Technology stakeholders.",
    "Conducted user interviews and mapped user journeys to understand needs, pain points, process gaps, and improvement opportunities.",
    "Explained concepts, options, and trade-offs to diverse audiences, helping business and technology teams align around practical change decisions.",
    "Built business cases for process and technology improvements, estimating hard benefits such as ROI and cost savings and soft benefits such as performance gains.",
    "Mentored team members, removed blockers, supported delivery discipline, and influenced team direction through structured communication and governance.",
]:
    add_bullet(doc, item)

add_heading(doc, "Selected Product / Technical Evidence")
for item in [
    "OnePlan / Board SaaS transformation: contributed to implementation of a large SaaS-enabled planning solution with Board at the core, balancing roadmap, governance, cost, stakeholders, and delivery risk.",
    "AI-assisted CRM prototype: built a Python-based CRM application for a real business user to manage clients, appointments, finance, and operational workflows; iterated with the user using AI-assisted development tools.",
    "Architecture pragmatism: challenge unnecessary service fragmentation where it increases maintenance cost without clear business value; favor right-sized architecture that is easier to evolve, operate, and govern.",
]:
    add_bullet(doc, item)

add_heading(doc, "Tools and Technology")
add_inline_section(doc, "Product / delivery: ", "Aha!, Jira, MS Project, ServiceNow, agile delivery, release governance, change management, backlog management, roadmap planning.")
add_inline_section(doc, "Technical fluency: ", "APIs, integrations, Azure, SQL, Python, PowerShell, Splunk, CI/CD concepts, data modeling, ETL coordination, GitHub/GitLab, PyCharm, VS Code.")
add_inline_section(doc, "Architecture / process: ", "Lucidchart, MS Visio, process mapping, requirements documentation, non-functional requirements, disaster recovery planning, support and governance documentation.")
add_inline_section(doc, "Governance / enterprise systems: ", "SOX, PII protection, Archer, SailPoint, CyberArk, Veeva, SOLMAN, Workday, Concur, Fieldglass.")
add_inline_section(doc, "Collaboration: ", "MS Teams, Zoom, Slack, Confluence; stakeholder engagement from operational teams to senior leadership.")

add_heading(doc, "Certifications")
for item in [
    "Registered Product Owner Certificate",
    "ITIL Foundation Certificate in IT Service Management v4",
    "Association of Chartered Certified Accountants",
]:
    add_bullet(doc, item)

add_heading(doc, "Education")
for item in [
    "Postgraduate Degree in Computer Science - IT Academy STEP, Kyiv, Ukraine, 2023",
    "Bachelor of Science in Applied Accounting - Oxford Brookes University, Oxford, UK, 2010",
    "Bachelor of Law - Kyiv National University, Kyiv, Ukraine, 2006",
]:
    add_bullet(doc, item)

doc.save(OUT)
print(OUT.resolve())
